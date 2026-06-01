import argparse
import json
import logging
import math
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, Dict, Any, List

import pandas as pd
import fitz  # PyMuPDF

from scripts import app_settings

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_WB = ROOT / "inventory_2026-04-17.xlsx"
INGEST_DIR = ROOT / "data" / "ingested"

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")


def excel_sheet_has_content(df: pd.DataFrame) -> bool:
    """Есть ли на листе хоть какой-то непустой текст (ячейки или осмысленные имена столбцов)."""
    if df is None:
        return False
    d = df.fillna("")
    for _, row in d.iterrows():
        for v in row:
            if str(v).strip():
                return True
    for c in d.columns:
        sc = str(c).strip()
        if sc and not sc.lower().startswith("unnamed"):
            return True
    return False


def _excel_cell_to_md(value: object) -> str:
    """Одна ячейка → текст markdown-таблицы (без переносов и сырых |)."""
    s = re.sub(r"[\r\n\u2028\u2029\v\f]+", " ", str(value))
    return s.replace("|", r"\|").strip()


def read_excel_workbook_sheets(file_path: Path) -> Dict[str, pd.DataFrame]:
    """Все листы .xlsx / .xls в порядке вкладок (явный ExcelFile — надёжнее, чем полагаться на edge-кейсы read_excel)."""
    ext = file_path.suffix.lower()
    if ext == ".xlsx":
        xl = pd.ExcelFile(file_path, engine="openpyxl")
        return {name: xl.parse(name) for name in xl.sheet_names}
    if ext == ".xls":
        last_err: Optional[BaseException] = None
        for eng in ("xlrd", "calamine", None):
            try:
                xl = pd.ExcelFile(file_path, engine=eng) if eng is not None else pd.ExcelFile(file_path)
                return {name: xl.parse(name) for name in xl.sheet_names}
            except Exception as e:
                last_err = e
        raise RuntimeError(f"Не удалось прочитать .xls ({file_path.name}): {last_err}") from last_err
    raise ValueError(f"read_excel_workbook_sheets: ожидался .xls или .xlsx, получено {ext}")


def excel_or_csv_workbook_to_text_pages(file_path: Path) -> List[str]:
    """Текстовые «страницы» для pages.jsonl: один непустой лист = одна markdown-таблица (одна строка jsonl).

    Порядок — порядок вкладок; пустые листы пропускаются. Дробление листа на несколько страниц
    только если задано ``ALTIORA_EXCEL_SHEET_SPLIT_CHARS`` > 0 (страховка для гигантских таблиц).
    """
    ext = file_path.suffix.lower()
    try:
        split_chars = int(os.getenv("ALTIORA_EXCEL_SHEET_SPLIT_CHARS", "0").strip() or "0")
    except ValueError:
        split_chars = 0
    split_chars = max(0, split_chars)

    pages: List[str] = []
    if ext == ".csv":
        sheets: Dict[str, pd.DataFrame] = {"Sheet1": pd.read_csv(file_path)}
    elif ext in (".xls", ".xlsx"):
        sheets = read_excel_workbook_sheets(file_path)
    else:
        raise ValueError(f"excel_or_csv_workbook_to_text_pages: не табличный формат {ext}")

    for sheet_name, df in sheets.items():
        if not excel_sheet_has_content(df):
            continue
        df = df.fillna("")
        headers = df.columns.astype(str).tolist()
        header_str = "| " + " | ".join([_excel_cell_to_md(h) for h in headers]) + " |"
        sep_str = "| " + " | ".join(["---"] * len(headers)) + " |"

        row_strs: List[str] = []
        for _, row in df.iterrows():
            row_strs.append(
                "| " + " | ".join([_excel_cell_to_md(x) for x in row]) + " |"
            )

        if not row_strs and df.shape[0] == 0:
            md_table = "\n".join([header_str, sep_str])
            pages.append(f"# Лист: {sheet_name}\n\n[Таблица]\n{md_table}")
            continue

        if split_chars <= 0:
            md_table = "\n".join([header_str, sep_str] + row_strs)
            pages.append(f"# Лист: {sheet_name}\n\n[Таблица]\n{md_table}")
            continue

        row_groups: List[List[str]] = []
        current_table_rows: List[str] = []
        current_len = len(header_str) + len(sep_str) + 2
        for row_str in row_strs:
            row_len = len(row_str) + 1
            if current_len + row_len > split_chars and current_table_rows:
                row_groups.append(current_table_rows)
                current_table_rows = []
                current_len = len(header_str) + len(sep_str) + 2
            current_table_rows.append(row_str)
            current_len += row_len
        if current_table_rows:
            row_groups.append(current_table_rows)

        nfrag = len(row_groups)
        for i, rows in enumerate(row_groups, start=1):
            md_table = "\n".join([header_str, sep_str] + rows)
            title = sheet_name if nfrag == 1 else f"{sheet_name} (часть {i}/{nfrag})"
            pages.append(f"# Лист: {title}\n\n[Таблица]\n{md_table}")

    return pages


def patch_manifest_metadata(manifest_path: Path, **extra: object) -> None:
    """Обновляет manifest.json после извлечения (без падения пайплайна)."""
    if not manifest_path.exists():
        return
    try:
        with manifest_path.open("r", encoding="utf-8") as f:
            m = json.load(f)
        m.update(extra)
        with manifest_path.open("w", encoding="utf-8") as f:
            json.dump(m, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logging.warning(f"Не удалось обновить manifest {manifest_path}: {e}")


def _running_strip_enabled() -> bool:
    v = (os.environ.get("ALTIORA_STRIP_RUNNING_LINES") or "1").strip().lower()
    return v not in ("0", "false", "no", "off")


def _line_key(line: str) -> str:
    return " ".join(line.split()).lower()


def strip_running_headers_footers(
    pages: List[str],
    *,
    top_scan: int = app_settings.RUNNING_TOP_SCAN,
    bottom_scan: int = app_settings.RUNNING_BOTTOM_SCAN,
    min_fraction: float = app_settings.RUNNING_MIN_FRACTION,
    max_line_len: int = app_settings.RUNNING_MAX_LINE_LEN,
    min_line_len: int = app_settings.RUNNING_MIN_LINE_LEN,
) -> List[str]:
    """Удаляет колонтитулы: строки, которые на **≥ min_fraction** страницах (по умолчанию 90 %) стоят
    **первой** или **последней** непустой строкой.

    Порог по доле страниц: ``ALTIORA_RUNNING_LINES_MIN_FRACTION`` (например ``0.92``), иначе аргумент
    ``min_fraction`` (по умолчанию 0.9). Отключение снятия: ``ALTIORA_STRIP_RUNNING_LINES=0``.
    """
    if not pages or not _running_strip_enabled():
        return list(pages)
    n = len(pages)
    if n < 3:
        return list(pages)

    mf = max(0.5, min(1.0, min_fraction))

    # Не менее ceil(mf * n) страниц с той же первой/последней строкой (для mf=0.9 — не ниже 90 %).
    need = max(1, math.ceil(mf * n - 1e-12))

    def _first_nonempty(s: str) -> str:
        for ln in s.splitlines():
            t = ln.strip()
            if t:
                return t
        return ""

    def _last_nonempty(s: str) -> str:
        for ln in reversed(s.splitlines()):
            t = ln.strip()
            if t:
                return t
        return ""

    def _strip_first_line_matching(s: str, key: str) -> str:
        lines = s.splitlines()
        i = 0
        while i < len(lines):
            t = lines[i].strip()
            if not t:
                i += 1
                continue
            if len(t) < min_line_len or len(t) > max_line_len:
                return s
            if _line_key(t) == key:
                i += 1
                return "\n".join(lines[i:]).strip()
            return s
        return s

    def _strip_last_line_matching(s: str, key: str) -> str:
        lines = s.splitlines()
        j = len(lines) - 1
        while j >= 0:
            t = lines[j].strip()
            if not t:
                j -= 1
                continue
            if len(t) < min_line_len or len(t) > max_line_len:
                return s
            if _line_key(t) == key:
                return "\n".join(lines[:j]).strip()
            return s
        return s

    out = list(pages)

    for _ in range(top_scan):
        cnt: Counter[str] = Counter()
        for raw in out:
            fn = _first_nonempty(raw)
            if fn and min_line_len <= len(fn) <= max_line_len:
                cnt[_line_key(fn)] += 1
        if not cnt:
            break
        best_k, best_c = cnt.most_common(1)[0]
        if best_c < need or not best_k:
            break
        out = [_strip_first_line_matching(p, best_k) for p in out]

    for _ in range(bottom_scan):
        cnt = Counter()
        for raw in out:
            ln = _last_nonempty(raw)
            if ln and min_line_len <= len(ln) <= max_line_len:
                cnt[_line_key(ln)] += 1
        if not cnt:
            break
        best_k, best_c = cnt.most_common(1)[0]
        if best_c < need or not best_k:
            break
        out = [_strip_last_line_matching(p, best_k) for p in out]

    return out


def resolve_soffice_path() -> Optional[str]:
    """Путь к LibreOffice headless (Office → PDF).

    Порядок: ``ALTIORA_SOFFICE`` / ``SOFFICE_PATH`` → ``LIBREOFFICE_HOME`` → ``PATH`` →
    типичные каталоги Windows (в т.ч. ``LibreOffice 24.x\\program\\soffice.exe``) → ``/usr/bin/soffice``.
    """
    env = os.environ

    def _ok(p: str) -> Optional[str]:
        p = (p or "").strip().strip('"')
        return p if p and os.path.isfile(p) else None

    for key in ("ALTIORA_SOFFICE", "SOFFICE_PATH"):
        found = _ok(env.get(key, ""))
        if found:
            return found

    lo_home = env.get("LIBREOFFICE_HOME", "").strip().strip('"')
    if lo_home:
        root = Path(lo_home)
        for rel in (Path("program") / "soffice.exe", Path("program") / "soffice"):
            cand = str(root / rel)
            if os.path.isfile(cand):
                return cand

    for candidate in (
        shutil.which("soffice"),
        shutil.which("soffice.exe"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
    ):
        found = _ok(candidate or "")
        if found:
            return found

    if os.name == "nt":
        for pf_key in ("ProgramFiles", "ProgramFiles(x86)"):
            base = env.get(pf_key) or ""
            if not base:
                continue
            base_p = Path(base)
            if not base_p.is_dir():
                continue
            matches: list[Path] = []
            try:
                matches = [p for p in base_p.glob("LibreOffice*") if p.is_dir()]
            except OSError:
                matches = []
            for lo_dir in sorted(matches, key=lambda p: p.name.lower(), reverse=True):
                exe = lo_dir / "program" / "soffice.exe"
                if exe.is_file():
                    return str(exe)

    return None


def convert_office_source_to_pdf(src: Path, work_dir: Path) -> Path:
    """Копирует файл в work_dir, конвертирует в PDF через LibreOffice; возвращает путь к PDF."""
    work_dir.mkdir(parents=True, exist_ok=True)
    soffice = resolve_soffice_path()
    if not soffice:
        raise RuntimeError("Не найден LibreOffice (soffice) для конвертации Office → PDF.")
    local_copy = work_dir / f"lo_source{src.suffix.lower()}"
    shutil.copy2(src, local_copy)
    cmd = [str(soffice), "--headless", "--convert-to", "pdf", "--outdir", str(work_dir), str(local_copy)]
    subprocess.run(cmd, check=True, capture_output=True)
    pdf_path = work_dir / f"{local_copy.stem}.pdf"
    if not pdf_path.is_file():
        raise RuntimeError("LibreOffice не создал PDF после конвертации.")
    return pdf_path


def write_api_native_ingest_no_soffice(
    source_id: str,
    source_file: Path,
    extension: str,
    out_root: Path = INGEST_DIR,
) -> tuple[int, str]:
    """Пишет ``pages.jsonl`` + ``manifest`` без LibreOffice и без YOLO (одна «страница»).

    Поддержка: ``.txt``, ``.docx``, ``.rtf``. Текст из docx — параграфы и простые таблицы
    (без точного соответствия вёрстке PDF). Пустой документ → ``(0, parser_name)``.
    """
    started = time.perf_counter()
    ext = extension.lower()
    parser_name = "api_native_unknown"
    text = ""

    if ext == ".txt":
        parser_name = "api_native_txt"
        text = source_file.read_text(encoding="utf-8", errors="replace")
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
    elif ext == ".rtf":
        parser_name = "api_native_rtf"
        from striprtf.striprtf import rtf_to_text

        raw = source_file.read_text(encoding="utf-8", errors="surrogateescape")
        text = rtf_to_text(raw)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
    elif ext == ".docx":
        parser_name = "api_native_docx"
        import docx

        document = docx.Document(source_file)
        parts: List[str] = []
        for p in document.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)
        for table in document.tables:
            rows_md: List[str] = []
            for row in table.rows:
                cells = [c.text.replace("\n", " ").replace("|", "\\|") for c in row.cells]
                if any(x.strip() for x in cells):
                    rows_md.append("| " + " | ".join(cells) + " |")
            if rows_md:
                parts.append("\n".join(rows_md))
        text = re.sub(r"\n{3,}", "\n\n", "\n\n".join(parts)).strip()
    else:
        raise ValueError(f"write_api_native_ingest_no_soffice: неподдерживаемое расширение {ext}")

    out_dir = out_root / source_id
    out_dir.mkdir(parents=True, exist_ok=True)

    if not text:
        manifest = {
            "source_id": source_id,
            "file_path": str(source_file.resolve()).replace("\\", "/"),
            "format": ext.lstrip("."),
            "status": "error",
            "page_count": 0,
            "error": "Пустой текст после нативного разбора (без LibreOffice).",
            "notes": [parser_name],
            "route": "native_office_api",
            "ocr_mode": "off",
            "parser": parser_name,
            "process_time_sec": round(time.perf_counter() - started, 3),
        }
        with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)
        return 0, parser_name

    try:
        rel_fp = str(source_file.resolve().relative_to(ROOT)).replace("\\", "/")
    except ValueError:
        rel_fp = str(source_file.resolve()).replace("\\", "/")

    with (out_dir / "pages.jsonl").open("w", encoding="utf-8") as f:
        rec = {"source_id": source_id, "page": 1, "text": text, "parser": parser_name}
        f.write(json.dumps(rec, ensure_ascii=False) + "\n")

    manifest = {
        "source_id": source_id,
        "file_path": rel_fp,
        "format": ext.lstrip("."),
        "status": "ok",
        "page_count": 1,
        "error": None,
        "notes": [
            f"{parser_name}: без LibreOffice/YOLO. Для распознавания как у отсканированного PDF "
            "установите LibreOffice (конвертация в PDF → layout + OCR)."
        ],
        "route": "native_office_api",
        "ocr_mode": "off",
        "parser": parser_name,
        "process_time_sec": round(time.perf_counter() - started, 3),
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)

    return 1, parser_name


@dataclass
class DocProfile:
    source_id: str
    file_path: Path
    extension: str
    file_size_mb: float
    is_image: bool
    is_pdf: bool
    is_office: bool
    page_count: int
    text_ratio: float
    has_text_layer: bool

def detect_document_profile(source_id: str, file_path: Path) -> DocProfile:
    """Определяет профиль документа дешево и быстро (без ML-моделей)."""
    ext = file_path.suffix.lower()
    size_mb = file_path.stat().st_size / (1024 * 1024)
    is_image = ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif")
    is_pdf = ext == ".pdf"
    is_office = ext in (".docx", ".doc", ".rtf", ".xls", ".xlsx", ".txt", ".csv")
    
    page_count = 1
    text_ratio = 0.0
    has_text_layer = False
    
    if is_pdf:
        try:
            with fitz.open(file_path) as doc:
                page_count = len(doc)
                sample_pages = min(5, page_count)
                total_chars = 0
                for i in range(sample_pages):
                    text = doc[i].get_text()
                    total_chars += len(text.strip())
                text_ratio = total_chars / sample_pages if sample_pages > 0 else 0
                has_text_layer = text_ratio > 50
        except Exception as e:
            logging.warning(f"[{source_id}] Не удалось проанализировать PDF через PyMuPDF: {e}")
    elif is_office:
        has_text_layer = True
    elif is_image:
        has_text_layer = False
        
    return DocProfile(
        source_id=source_id,
        file_path=file_path,
        extension=ext,
        file_size_mb=size_mb,
        is_image=is_image,
        is_pdf=is_pdf,
        is_office=is_office,
        page_count=page_count,
        text_ratio=text_ratio,
        has_text_layer=has_text_layer
    )

def decide_route(profile: DocProfile, args: argparse.Namespace) -> Dict[str, str]:
    """Принимает решение о маршрутизации."""
    parser_mode = getattr(args, "parser", "auto")
    if parser_mode == "docling":
        return {"route": "docling", "reason": "Явно выбран parser=docling"}
    if parser_mode == "yolo":
        return {"route": "heavy", "reason": "Явно выбран parser=yolo"}

    if getattr(args, 'force_heavy', False):
        return {"route": "heavy", "reason": "Принудительный запуск Heavy ветки (YOLO)"}
    if args.force_light:
        return {"route": "light", "reason": "Установлен флаг --force-light"}
        
    # Автоматическая маршрутизация
    if profile.is_image:
        return {"route": "heavy", "reason": "Формат изображения, требуется Layout-анализ и OCR"}
        
    if profile.is_pdf:
        # Всегда heavy для PDF: единый пайплайн с таблицами/формулами (см. ingest_yolo_doctr).
        return {
            "route": "heavy",
            "reason": (
                f"PDF — маршрут Heavy (YOLO+OCR); текстовый слой ~{profile.text_ratio:.1f} симв/стр "
                "(легкий PyMuPDF только с --force-light)."
            ),
        }
            
    if profile.is_office:
        ext = profile.extension.lower()
        if ext in (".doc", ".docx", ".rtf"):
            return {
                "route": "heavy",
                "reason": "Word/RTF → PDF (LibreOffice) → YOLO — тот же принцип, что в API.",
            }
        return {
            "route": "light",
            "reason": "Excel/CSV/TXT — в батче нативный разбор (markdown-таблицы); YOLO в ingest_yolo_doctr только по PDF.",
        }
        
    return {"route": "heavy", "reason": "Резервный маршрут по умолчанию (Heavy)"}


class Extractor:
    def extract(self, profile: DocProfile, **kwargs) -> int:
        raise NotImplementedError()

class LightExtractor(Extractor):
    def __init__(self, output_root: Path = INGEST_DIR):
        self.output_root = output_root

    def extract(self, profile: DocProfile, **kwargs) -> int:
        started = time.perf_counter()
        logging.info(f"[{profile.source_id}] Запуск Light Extractor...")
        pages_content = {}
        out_dir = self.output_root / profile.source_id
        out_dir.mkdir(parents=True, exist_ok=True)
        parser_name = "light_unknown"
        error_msg = None
        
        try:
            if profile.is_pdf:
                parser_name = "light_pymupdf"
                with fitz.open(profile.file_path) as doc:
                    for i, page in enumerate(doc):
                        text = page.get_text()
                        text = re.sub(r'\n{3,}', '\n\n', text).strip()
                        if text:
                            pages_content[i+1] = text
                if len(pages_content) >= 3:
                    keys = sorted(pages_content.keys())
                    texts = strip_running_headers_footers([pages_content[k] for k in keys])
                    for k, t in zip(keys, texts):
                        pages_content[k] = t if t.strip() else pages_content[k]
            elif profile.extension == ".docx":
                parser_name = "light_docx"
                import docx
                doc = docx.Document(profile.file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
                text = re.sub(r'\n{3,}', '\n\n', text).strip()
                if text:
                    pages_content[1] = text
            elif profile.extension == ".rtf":
                parser_name = "light_rtf"
                from striprtf.striprtf import rtf_to_text
                with open(profile.file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = rtf_to_text(f.read())
                text = re.sub(r'\n{3,}', '\n\n', text).strip()
                if text:
                    pages_content[1] = text
            elif profile.extension in (".xls", ".xlsx", ".csv"):
                parser_name = "light_pandas"
                for i, text in enumerate(excel_or_csv_workbook_to_text_pages(profile.file_path), start=1):
                    pages_content[i] = text
            elif profile.extension in (".doc", ".txt"):
                parser_name = f"light_{profile.extension.lstrip('.')}_libreoffice"
                with tempfile.TemporaryDirectory() as tmpdirname:
                    soffice = resolve_soffice_path()
                    if not soffice:
                        raise RuntimeError("Не найден LibreOffice (soffice) для конвертации .doc")
                    cmd = [str(soffice), "--headless", "--convert-to", "pdf", "--outdir", tmpdirname, str(profile.file_path)]
                    subprocess.run(cmd, check=True, capture_output=True)
                    pdf_path = Path(tmpdirname) / (profile.file_path.stem + ".pdf")
                    if not pdf_path.exists():
                        raise RuntimeError("Не удалось конвертировать .doc в PDF через LibreOffice.")
                    with fitz.open(pdf_path) as doc:
                        for i, page in enumerate(doc):
                            text = page.get_text()
                            text = re.sub(r'\n{3,}', '\n\n', text).strip()
                            if text:
                                pages_content[i+1] = text
                    if len(pages_content) >= 3:
                        keys = sorted(pages_content.keys())
                        texts = strip_running_headers_footers([pages_content[k] for k in keys])
                        for k, t in zip(keys, texts):
                            pages_content[k] = t if t.strip() else pages_content[k]
            else:
                raise ValueError(f"Неподдерживаемый формат для Light ветки: {profile.extension}")
        except Exception as e:
            error_msg = str(e)
            logging.error(f"[{profile.source_id}] Ошибка Light ветки: {e}")
            
        written_pages = 0
        if not error_msg:
            pages_jsonl = out_dir / "pages.jsonl"
            with pages_jsonl.open("w", encoding="utf-8") as f:
                for p_num in sorted(pages_content.keys()):
                    record = {
                        "source_id": profile.source_id,
                        "page": p_num,
                        "text": pages_content[p_num],
                        "parser": parser_name
                    }
                    f.write(json.dumps(record, ensure_ascii=False) + "\n")
                    written_pages += 1
                    
        manifest = {
            "source_id": profile.source_id,
            "file_path": str(profile.file_path.relative_to(ROOT)).replace("\\", "/"),
            "format": profile.extension.lstrip('.'),
            "status": "ok" if written_pages > 0 else "error",
            "page_count": written_pages,
            "error": error_msg,
            "notes": [f"Parsed via LightExtractor ({parser_name})"],
            "route": "light",
            "ocr_mode": "off",
            "parser": parser_name,
            "process_time_sec": round(time.perf_counter() - started, 3),
        }
        with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)
            
        return written_pages

class HeavyExtractor(Extractor):
    def __init__(self, isolate: bool, workbook: Path, output_root: Path = INGEST_DIR):
        self.isolate = isolate
        self.workbook = workbook
        self.output_root = output_root

    def extract(self, profile: DocProfile) -> int:
        started = time.perf_counter()
        ext = profile.extension.lower()
        work_dir = self.output_root / profile.source_id

        # Как в API: xlsx/csv — pandas, не растр YOLO.
        if ext in (".xls", ".xlsx", ".csv"):
            le = LightExtractor(self.output_root)
            pages = le.extract(profile)
            manifest_path = work_dir / "manifest.json"
            elapsed = round(time.perf_counter() - started, 3)
            patch_manifest_metadata(
                manifest_path,
                route="native_table",
                parser="light_pandas",
                ocr_mode="off",
                process_time_sec=elapsed,
                notes=["Табличный файл: pandas → markdown (аналог ветки native_table в API)."],
            )
            return pages

        # Как в API: doc/docx/rtf → PDF (LibreOffice) → YOLO.
        if ext in (".doc", ".docx", ".rtf"):
            try:
                pdf_path = convert_office_source_to_pdf(profile.file_path, work_dir)
            except Exception as e:
                logging.error(f"[{profile.source_id}] Конвертация Office→PDF: {e}")
                return 0
            try:
                from scripts.ingest_yolo_doctr import YoloDoctrExtractor

                yolo = YoloDoctrExtractor()
                pages = yolo.extract(profile.source_id, pdf_path, self.output_root)
            except Exception as e:
                logging.error(f"[{profile.source_id}] YOLO после конвертации: {e}")
                return 0
            manifest_path = work_dir / "manifest.json"
            elapsed = round(time.perf_counter() - started, 3)
            orig_fp = str(profile.file_path.relative_to(ROOT)).replace("\\", "/")
            patch_manifest_metadata(
                manifest_path,
                process_time_sec=elapsed,
                route="heavy",
                parser="yolo_doctr",
                ocr_mode="on",
                file_path=orig_fp,
                notes=[
                    "Parsed via DocLayout-YOLO + Native Text/EasyOCR (Hybrid); "
                    "исходник Word/RTF сконвертирован в PDF через LibreOffice."
                ],
            )
            return pages

        cmd = [
            sys.executable,
            str(ROOT / "scripts" / "ingest_yolo_doctr.py"),
            "--workbook",
            str(self.workbook),
            "--only",
            profile.source_id,
            "--output-root",
            str(self.output_root),
        ]

        logging.info(f"[{profile.source_id}] Запуск Heavy Extractor (YOLO): {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=False)

        if result.returncode != 0:
            logging.error(f"[{profile.source_id}] Ошибка subprocess (код {result.returncode})")
            return 0

        manifest_path = work_dir / "manifest.json"
        pages = 0
        if manifest_path.exists():
            try:
                with open(manifest_path, "r", encoding="utf-8") as f:
                    man = json.load(f)
                pages = man.get("page_count", 0)
            except Exception:
                pass
        elapsed = round(time.perf_counter() - started, 3)
        patch_manifest_metadata(
            manifest_path,
            parser="yolo_doctr",
            route="heavy",
            process_time_sec=elapsed,
        )

        return pages


class DoclingExtractor(Extractor):
    def __init__(self, output_root: Path = INGEST_DIR):
        self.output_root = output_root
        self.last_returncode: Optional[int] = None

    def extract(self, profile: DocProfile) -> int:
        started = time.perf_counter()
        self.last_returncode = None
        cmd = [
            sys.executable,
            str(ROOT / "scripts" / "ingest_docling.py"),
            "--direct-file",
            str(profile.file_path.resolve()),
            "--only",
            profile.source_id,
            "--output-root",
            str(self.output_root),
            "--resume",
        ]
        logging.info(f"[{profile.source_id}] Запуск Docling Extractor: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=False, cwd=str(ROOT))
        self.last_returncode = result.returncode
        if result.returncode != 0:
            logging.error(f"[{profile.source_id}] Ошибка Docling subprocess (код {result.returncode})")
            return 0

        manifest_path = self.output_root / profile.source_id / "manifest.json"
        pages = 0
        if manifest_path.exists():
            try:
                with manifest_path.open("r", encoding="utf-8") as f:
                    man = json.load(f)
                pages = int(man.get("page_count", 0))
            except Exception:
                pass
        elapsed = round(time.perf_counter() - started, 3)
        patch_manifest_metadata(
            manifest_path,
            parser="docling",
            route="docling",
            process_time_sec=elapsed,
        )
        return pages

def main():
    parser = argparse.ArgumentParser(description="Универсальный Ingest Router для документов (маршрутизация на Light / Heavy).")
    parser.add_argument("--workbook", type=Path, default=DEFAULT_WB, help="Путь к Excel реестру")
    parser.add_argument("--only", type=str, default="", help="SRC-id через запятую")
    parser.add_argument("--resume", action="store_true", help="Пропускать уже обработанные документы")
    parser.add_argument(
        "--parser",
        choices=("auto", "yolo", "docling"),
        default="auto",
        help="auto: PDF/изображения/Word/RTF → Heavy; Excel/CSV/TXT → нативно; yolo/docling — явный выбор",
    )
    parser.add_argument("--force-light", action="store_true", help="Принудительно направить все в Light ветку")
    parser.add_argument("--force-heavy", action="store_true", help="Принудительно направить все в Heavy ветку (YOLO)")
    parser.add_argument("--isolate", action="store_true", help="Запускать Heavy ветку в отдельном процессе (всегда активно для YOLO)")
    parser.add_argument("--output-root", type=Path, default=INGEST_DIR, help="Каталог для data/ingested результатов")
    
    args = parser.parse_args()

    if not args.workbook.is_file():
        logging.error(f"Файл реестра не найден: {args.workbook}")
        sys.exit(1)
        
    df = pd.read_excel(args.workbook, sheet_name="source_registry")
    only_ids = {s.strip() for s in args.only.split(",") if s.strip()}
    
    output_root = args.output_root
    output_root.mkdir(parents=True, exist_ok=True)
    heavy_extractor = HeavyExtractor(isolate=args.isolate, workbook=args.workbook, output_root=output_root)
    light_extractor = LightExtractor(output_root=output_root)
    docling_extractor = DoclingExtractor(output_root=output_root)
    
    processed = 0
    for _, row in df.iterrows():
        src = str(row.get("source_id", "")).strip()
        fp = str(row.get("file_path", "")).strip()
        fmt = str(row.get("format", "")).strip().lower()
        
        if not src or src == "nan" or not fp:
            continue
        if only_ids and src not in only_ids:
            continue
            
        allowed_formats = ("pdf", "docx", "doc", "rtf", "xls", "xlsx", "txt", "csv", "jpg", "jpeg", "png", "bmp", "tiff", "tif")
        if fmt not in allowed_formats:
            logging.warning(f"[{src}] Формат {fmt} не поддерживается")
            continue
            
        if args.resume:
            manifest_path = output_root / src / "manifest.json"
            if manifest_path.is_file():
                try:
                    with open(manifest_path, "r", encoding="utf-8") as f:
                        man = json.load(f)
                        if man.get("status") in ("ok", "empty"):
                            logging.info(f"[{src}] Уже обработан, пропускаем (--resume)...")
                            continue
                except Exception:
                    pass
                    
        abs_path = (ROOT / fp).resolve()
        if not abs_path.is_file():
            logging.warning(f"[{src}] Файл не найден: {abs_path}")
            continue
            
        profile = detect_document_profile(src, abs_path)
        decision = decide_route(profile, args)
        
        logging.info(f"[{src}] === Маршрутизация ===")
        logging.info(f"  Файл: {profile.file_path.name} ({profile.file_size_mb:.1f} MB)")
        logging.info(f"  Профиль: PDF={profile.is_pdf}, ТекстСлой={profile.has_text_layer}, Image={profile.is_image}")
        logging.info(f"  Выбран маршрут: {decision['route'].upper()}")
        logging.info(f"  Причина: {decision['reason']}")
        
        if decision["route"] == "light":
            pages = light_extractor.extract(profile)
        elif decision["route"] == "docling":
            pages = docling_extractor.extract(profile)
        else:
            pages = heavy_extractor.extract(profile)
            
        if pages > 0:
            logging.info(f"[{src}] Успешно завершено: {pages} страниц.")
            processed += 1
        else:
            logging.warning(f"[{src}] Документ извлечен с ошибкой или пуст.")
            
    logging.info(f"Универсальный Ingest завершен. Обработано документов: {processed}")

if __name__ == "__main__":
    main()